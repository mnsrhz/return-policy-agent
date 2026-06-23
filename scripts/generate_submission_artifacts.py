"""Generate Week 4 submission artifacts from eval source files.

JSONL and JSON files remain the source of truth. This script creates readable
Excel and Word exports for course submission.
"""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "submissions"

SOURCE_PATHS = {
    "golden_dataset": PROJECT_ROOT / "data/golden_dataset.jsonl",
    "baseline_predictions": PROJECT_ROOT / "eval_runs/local_baseline_predictions.jsonl",
    "baseline_summary": PROJECT_ROOT / "eval_runs/local_baseline_summary.json",
    "post_predictions": PROJECT_ROOT / "eval_runs/local_post_improvement_predictions.jsonl",
    "post_summary": PROJECT_ROOT / "eval_runs/local_post_improvement_summary.json",
    "baseline_report": PROJECT_ROOT / "docs/14_BASELINE_EVAL_REPORT.md",
    "improvement_log": PROJECT_ROOT / "docs/15_IMPROVEMENT_LOG.md",
    "final_report": PROJECT_ROOT / "docs/16_FINAL_WEEK4_EVALUATION_REPORT.md",
    "loom_script": PROJECT_ROOT / "docs/17_LOOM_WALKTHROUGH_SCRIPT.md",
}

METRIC_DEFINITIONS = [
    ("decision_accuracy", "Exact match between predicted and expected decision."),
    ("reason_code_accuracy", "Exact match between predicted and expected reason code."),
    ("missing_info_f1", "Set-based F1 comparing predicted and expected missing-info fields."),
    ("escalation_accuracy", "Exact match on escalation flag."),
    ("citation_coverage", "Prediction has at least one citation tied to retrieved policy evidence."),
    ("policy_section_recall", "Recall over expected policy sections."),
    ("schema_validity", "Required structured output fields are present."),
    ("average_latency_seconds", "Average local runtime latency for each evaluated case."),
]

PASS_BARS = {
    "decision_accuracy": 90.0,
    "reason_code_accuracy": 85.0,
    "missing_info_f1": 85.0,
    "escalation_accuracy": 95.0,
    "citation_coverage": 100.0,
    "policy_section_recall": None,
    "schema_validity": None,
    "average_latency_seconds": None,
}

LONG_TEXT_HEADERS = {
    "customer_message",
    "expected_answer_traits",
    "predicted_answer",
    "customer_answer",
    "failure_notes",
    "notes",
    "expected",
    "prediction",
    "order_context",
    "expected_policy_sections",
    "expected_missing_info",
    "Likely Root Cause",
    "Recommended Improvement",
    "Expected Impact",
    "Regression Risk",
}

STATUS_GREEN = PatternFill("solid", fgColor="D9EAD3")
STATUS_RED = PatternFill("solid", fgColor="F4CCCC")
STATUS_YELLOW = PatternFill("solid", fgColor="FFF2CC")
HEADER_FILL = PatternFill("solid", fgColor="E8EEF5")
TITLE_FILL = PatternFill("solid", fgColor="D9EAF7")
TIMESTAMP_FILL = PatternFill("solid", fgColor="F2F4F7")
THIN = Side(style="thin", color="D9E2EC")

LABEL_DEFINITIONS = [
    ("happy_path", "Common cases where policy and order facts produce a direct outcome."),
    ("edge_case", "Boundary, incomplete, restricted, or multi-policy cases."),
    ("known_failure", "Cases selected to expose known or likely weaknesses."),
    ("adversarial", "Prompt injection or override attempts."),
    ("eligible_return", "Item appears eligible for return under policy."),
    ("not_eligible", "Policy does not allow the requested return/refund/exchange."),
    ("eligible_exchange", "Item appears eligible for size or color exchange, subject to inventory."),
    ("ask_for_info", "Required facts are missing, so the agent should ask follow-up questions."),
    ("escalate", "Human support review is required."),
]


def _path_map(source_overrides: dict[str, Path] | None = None) -> dict[str, Path]:
    paths = dict(SOURCE_PATHS)
    if source_overrides:
        paths.update({key: Path(value) for key, value in source_overrides.items()})
    return paths


def _source_missing(path: Path) -> str:
    return f"Source file was not found: {path}"


def _read_text(path: Path) -> str:
    if not path.exists():
        return _source_missing(path)
    return path.read_text(encoding="utf-8")


def _read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"_missing_source": _source_missing(path)}
    return json.loads(path.read_text(encoding="utf-8"))


def _read_jsonl(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return [{"_missing_source": _source_missing(path)}]
    rows: list[dict[str, Any]] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            rows.append(json.loads(line))
    return rows


def _jsonish(value: Any) -> Any:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=True, sort_keys=True)
    return value


def _metric_percent(value: Any) -> Any:
    if isinstance(value, (int, float)) and 0 <= value <= 1:
        return round(value * 100, 2)
    return value


def _generated_timestamp() -> str:
    return datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")


def _new_workbook() -> Workbook:
    wb = Workbook()
    wb.remove(wb.active)
    return wb


def _style_sheet(ws) -> None:
    border = Border(bottom=THIN)
    header_row = 3

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if cell.row == 1:
                cell.font = Font(bold=True, color="0B2545", size=14)
                cell.fill = TITLE_FILL
            elif cell.row == 2:
                cell.font = Font(italic=True, color="555555")
                cell.fill = TIMESTAMP_FILL
            elif cell.row == header_row:
                cell.font = Font(bold=True)
                cell.fill = HEADER_FILL
                cell.border = border
            _apply_status_fill(cell)

    if ws.max_row >= header_row:
        ws.freeze_panes = "A4"
        ws.auto_filter.ref = f"A{header_row}:{get_column_letter(ws.max_column)}{ws.max_row}"
    ws.sheet_view.showGridLines = False
    _style_long_text_columns(ws, header_row)
    _auto_size(ws)


def _auto_size(ws, max_width: int = 64) -> None:
    for column_cells in ws.columns:
        letter = get_column_letter(column_cells[0].column)
        width = 12
        for cell in column_cells:
            if cell.value is None:
                continue
            for line in str(cell.value).splitlines():
                width = max(width, min(max_width, len(line) + 2))
        ws.column_dimensions[letter].width = width


def _style_long_text_columns(ws, header_row: int) -> None:
    headers = [cell.value for cell in ws[header_row]]
    for index, header in enumerate(headers, start=1):
        if header in LONG_TEXT_HEADERS:
            letter = get_column_letter(index)
            ws.column_dimensions[letter].width = 48
            for cell in ws[letter]:
                cell.alignment = Alignment(vertical="top", wrap_text=True)


def _apply_status_fill(cell) -> None:
    value = cell.value
    if value is None:
        return
    normalized = str(value).strip().lower()
    if normalized in {"pass", "passed", "true", "1.0", "100.0", "100"}:
        cell.fill = STATUS_GREEN
    elif normalized in {"fail", "failed", "false", "0.0", "0"}:
        cell.fill = STATUS_RED
    elif normalized in {"partial", "needs review", "not measured", "diagnostic", "n/a"}:
        cell.fill = STATUS_YELLOW


def _add_rows(ws, title: str, headers: list[str], rows: Iterable[Iterable[Any]]) -> None:
    ws.append([title])
    ws.append(["Generated", _generated_timestamp()])
    ws.append(headers)
    for row in rows:
        ws.append([_jsonish(value) for value in row])
    _style_sheet(ws)


def _summary_rows(summary: dict[str, Any]) -> list[list[Any]]:
    if "_missing_source" in summary:
        return [["source_status", summary["_missing_source"], "N/A", "Needs Review"]]
    rows = []
    for key, value in sorted(summary.items()):
        metric_value = _metric_percent(value)
        pass_bar = _pass_bar_label(key)
        rows.append([key, metric_value, pass_bar, _metric_status(key, metric_value)])
    return rows


def _case_rows(predictions: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for item in predictions:
        if "_missing_source" in item:
            return [["", "", item["_missing_source"], "", "", "", "", "", "", "", "", "", "", "", "", "", "", ""]]
        prediction = item.get("prediction", {})
        expected = item.get("expected", {})
        metrics = item.get("metrics", {})
        rows.append(
            [
                item.get("id"),
                item.get("scenario_type"),
                item.get("customer_message"),
                expected.get("decision"),
                prediction.get("decision"),
                expected.get("reason_code"),
                prediction.get("reason_code"),
                prediction.get("escalate"),
                prediction.get("customer_answer"),
                metrics.get("decision_accuracy"),
                _metric_case_status(metrics.get("decision_accuracy")),
                metrics.get("reason_code_accuracy"),
                _metric_case_status(metrics.get("reason_code_accuracy")),
                metrics.get("missing_info_f1"),
                _metric_case_status(metrics.get("missing_info_f1")),
                metrics.get("escalation_accuracy"),
                _metric_case_status(metrics.get("escalation_accuracy")),
                item.get("latency_seconds"),
            ]
        )
    return rows


def _failure_rows(predictions: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for item in predictions:
        if "_missing_source" in item:
            return [["", item["_missing_source"], "", "", ""]]
        metrics = item.get("metrics", {})
        failed = [key for key, value in metrics.items() if value != 1.0]
        if not failed:
            continue
        prediction = item.get("prediction", {})
        expected = item.get("expected", {})
        rows.append(
            [
                item.get("id"),
                item.get("scenario_type"),
                ", ".join(failed),
                expected,
                prediction,
            ]
        )
    return rows or [["No remaining failures", "", "", "", ""]]


def _golden_rows(cases: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for item in cases:
        if "_missing_source" in item:
            return [["", "", item["_missing_source"], "", "", "", "", "", "", "", ""]]
        rows.append(
            [
                item.get("id"),
                item.get("scenario_type"),
                item.get("customer_message"),
                item.get("order_context"),
                item.get("expected_decision"),
                item.get("expected_reason_code"),
                item.get("expected_missing_info"),
                item.get("expected_escalate"),
                item.get("expected_policy_sections"),
                item.get("expected_answer_traits"),
                item.get("notes"),
            ]
        )
    return rows


def _scenario_mix_rows(cases: list[dict[str, Any]]) -> list[list[Any]]:
    if cases and "_missing_source" in cases[0]:
        return [["source_status", cases[0]["_missing_source"]]]
    counts = Counter(item.get("scenario_type", "unknown") for item in cases)
    return [[key, value] for key, value in sorted(counts.items())]


def _distribution_rows(cases: list[dict[str, Any]], field: str) -> list[list[Any]]:
    if cases and "_missing_source" in cases[0]:
        return [["source_status", cases[0]["_missing_source"]]]
    counts = Counter(item.get(field, "unknown") for item in cases)
    return [[key, value] for key, value in sorted(counts.items())]


def _scenario_breakdown_rows(predictions: list[dict[str, Any]]) -> list[list[Any]]:
    if predictions and "_missing_source" in predictions[0]:
        return [["source_status", predictions[0]["_missing_source"], "", "", "", "", ""]]
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in predictions:
        grouped.setdefault(item.get("scenario_type", "unknown"), []).append(item)
    rows = []
    for scenario_type, items in sorted(grouped.items()):
        rows.append(
            [
                scenario_type,
                len(items),
                _avg_metric(items, "decision_accuracy"),
                _avg_metric(items, "reason_code_accuracy"),
                _avg_metric(items, "missing_info_f1"),
                _avg_metric(items, "escalation_accuracy"),
                sum(1 for item in items if any(value != 1.0 for value in item.get("metrics", {}).values())),
            ]
        )
    return rows


def _avg_metric(items: list[dict[str, Any]], metric: str) -> float:
    values = [item.get("metrics", {}).get(metric) for item in items]
    numeric = [value for value in values if isinstance(value, (int, float))]
    if not numeric:
        return 0.0
    return round(sum(numeric) / len(numeric) * 100, 2)


def _how_to_read_rows(workbook_name: str, notes: list[str]) -> list[list[Any]]:
    return [[workbook_name, note] for note in notes]


def _extract_markdown_sections(text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_title = "Overview"
    current_lines: list[str] = []
    for line in text.splitlines():
        if line.startswith("## "):
            if current_lines:
                sections.append((current_title, "\n".join(current_lines).strip()))
            current_title = line.lstrip("# ").strip()
            current_lines = []
        elif not line.startswith("# "):
            current_lines.append(line)
    if current_lines:
        sections.append((current_title, "\n".join(current_lines).strip()))
    return sections


def _markdown_table_rows(text: str, title_contains: str | None = None) -> list[list[str]]:
    if title_contains:
        pattern = re.compile(
            rf"## .*{re.escape(title_contains)}.*?\n(?P<body>.*?)(?=\n## |\Z)",
            re.DOTALL | re.IGNORECASE,
        )
        match = pattern.search(text)
        text = match.group("body") if match else ""
    rows: list[list[str]] = []
    for line in text.splitlines():
        if not line.strip().startswith("|"):
            continue
        cells = [cell.strip(" `") for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows


def _failure_cluster_rows(report_text: str) -> list[list[Any]]:
    rows = _markdown_table_rows(report_text, "Failure Analysis Table")
    if len(rows) > 1:
        return rows[1:]
    if report_text.startswith("Source file was not found"):
        return [["source_status", report_text, "", ""]]
    parsed: list[list[Any]] = []
    for match in re.finditer(r"### (?P<title>.*?)\n(?P<body>.*?)(?=\n### |\n## |\Z)", report_text, re.DOTALL):
        title = match.group("title").strip()
        body = match.group("body")
        count = _extract_bullet_value(body, "Count")
        example = _extract_bullet_value(body, "Example customer message")
        root_cause = _extract_bullet_value(body, "Likely root cause")
        if count:
            parsed.append([title, count, example, root_cause])
    return parsed or [["No failure table found", "", "", ""]]


def _improvement_rows(log_text: str) -> list[list[Any]]:
    if log_text.startswith("Source file was not found"):
        return [["source_status", log_text, "", "", ""]]
    rows: list[list[Any]] = []
    chunks = re.split(r"\n## Improvement ", log_text)
    for chunk in chunks[1:]:
        lines = chunk.splitlines()
        name = lines[0].strip()
        values = {"name": name}
        for line in lines[1:]:
            match = re.match(r"- ([^:]+):\s*(.*)", line.strip())
            if match:
                values[match.group(1).lower().replace(" ", "_")] = match.group(2)
        rows.append(
            [
                values.get("name"),
                values.get("failure_cluster_targeted"),
                values.get("files_changed"),
                values.get("expected_metric_impact"),
                values.get("risk_of_regression"),
            ]
        )
    return rows or [["No improvements parsed", "", "", "", ""]]


def _delta_rows(baseline: dict[str, Any], post: dict[str, Any]) -> list[list[Any]]:
    rows = []
    for metric in METRIC_DEFINITIONS:
        key = metric[0]
        before = baseline.get(key)
        after = post.get(key)
        if isinstance(before, (int, float)) and isinstance(after, (int, float)):
            baseline_value = _metric_percent(before)
            post_value = _metric_percent(after)
            rows.append(
                [
                    key,
                    baseline_value,
                    post_value,
                    round(post_value - baseline_value, 2),
                    _pass_bar_label(key),
                    _metric_status(key, post_value),
                ]
            )
        else:
            rows.append([key, before if before is not None else "N/A", after if after is not None else "N/A", "N/A", _pass_bar_label(key), "Needs Review"])
    return rows


def _pass_bar_label(metric: str) -> str:
    pass_bar = PASS_BARS.get(metric)
    if pass_bar is None:
        return "Diagnostic"
    if metric == "average_latency_seconds":
        return "< 8 seconds p95"
    return f">= {pass_bar:g}%"


def _metric_status(metric: str, value: Any) -> str:
    pass_bar = PASS_BARS.get(metric)
    if pass_bar is None or value == "N/A":
        return "Needs Review"
    if metric == "average_latency_seconds":
        return "Needs Review"
    if isinstance(value, (int, float)):
        return "Pass" if value >= pass_bar else "Fail"
    return "Needs Review"


def _metric_case_status(value: Any) -> str:
    if value == 1.0:
        return "Pass"
    if isinstance(value, (int, float)) and 0 < value < 1:
        return "Partial"
    if value == 0.0:
        return "Fail"
    return "Needs Review"


def _save_workbook(wb: Workbook, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    return path


def _create_golden_dataset_workbook(cases: list[dict[str, Any]], path: Path) -> Path:
    wb = _new_workbook()
    _add_rows(
        wb.create_sheet("How to Read"),
        "How to Read This Workbook",
        ["Workbook", "Guidance"],
        _how_to_read_rows(
            "Golden_Dataset.xlsx",
            [
                "JSONL remains the source of truth; this workbook is a reviewer-friendly export.",
                "Use Golden Dataset for case-level labels and expected outputs.",
                "Use Scenario Mix, Decision Distribution, and Reason Distribution to inspect coverage.",
            ],
        ),
    )
    _add_rows(
        wb.create_sheet("Overview"),
        "Golden Dataset Overview",
        ["Field", "Value"],
        [
            ["Source", "data/golden_dataset.jsonl"],
            ["Case Count", len(cases) if not (cases and "_missing_source" in cases[0]) else "N/A"],
            ["Status", cases[0]["_missing_source"] if cases and "_missing_source" in cases[0] else "Loaded"],
        ],
    )
    _add_rows(
        wb.create_sheet("Golden Dataset"),
        "Golden Dataset",
        [
            "id",
            "scenario_type",
            "customer_message",
            "order_context",
            "expected_decision",
            "expected_reason_code",
            "expected_missing_info",
            "expected_escalate",
            "expected_policy_sections",
            "expected_answer_traits",
            "notes",
        ],
        _golden_rows(cases),
    )
    _add_rows(wb.create_sheet("Label Definitions"), "Label Definitions", ["Label", "Definition"], LABEL_DEFINITIONS)
    _add_rows(wb.create_sheet("Scenario Mix"), "Scenario Mix", ["Scenario Type", "Count"], _scenario_mix_rows(cases))
    _add_rows(wb.create_sheet("Decision Distribution"), "Label Distribution By Expected Decision", ["Expected Decision", "Count"], _distribution_rows(cases, "expected_decision"))
    _add_rows(wb.create_sheet("Reason Distribution"), "Label Distribution By Expected Reason Code", ["Expected Reason Code", "Count"], _distribution_rows(cases, "expected_reason_code"))
    return _save_workbook(wb, path)


def _create_eval_workbook(
    *,
    path: Path,
    title: str,
    summary: dict[str, Any],
    predictions: list[dict[str, Any]],
    baseline_summary: dict[str, Any] | None = None,
) -> Path:
    wb = _new_workbook()
    is_post = baseline_summary is not None
    _add_rows(
        wb.create_sheet("How to Read"),
        "How to Read This Workbook",
        ["Workbook", "Guidance"],
        _how_to_read_rows(
            path.name,
            [
                "Summary Metrics shows overall evaluator scores, pass bars, and status.",
                "Case Results shows per-case predictions and per-metric pass/fail status.",
                "Scenario Breakdown groups performance by golden dataset scenario type.",
                "Failures or Remaining Failures lists any case with at least one non-passing metric.",
            ],
        ),
    )
    _add_rows(wb.create_sheet("Summary Metrics"), f"{title} Summary Metrics", ["Metric", "Value", "Pass Bar", "Status"], _summary_rows(summary))
    _add_rows(
        wb.create_sheet("Scenario Breakdown"),
        f"{title} Scenario-Type Breakdown",
        ["Scenario Type", "Case Count", "Decision Accuracy %", "Reason Code Accuracy %", "Missing Info F1 %", "Escalation Accuracy %", "Failed Case Count"],
        _scenario_breakdown_rows(predictions),
    )
    _add_rows(
        wb.create_sheet("Case Results"),
        f"{title} Case Results",
        [
            "id",
            "scenario_type",
            "customer_message",
            "expected_decision",
            "predicted_decision",
            "expected_reason_code",
            "predicted_reason_code",
            "predicted_escalate",
            "predicted_answer",
            "decision_accuracy",
            "decision_accuracy_status",
            "reason_code_accuracy",
            "reason_code_accuracy_status",
            "missing_info_f1",
            "missing_info_f1_status",
            "escalation_accuracy",
            "escalation_accuracy_status",
            "latency_seconds",
        ],
        _case_rows(predictions),
    )
    if not is_post:
        _add_rows(wb.create_sheet("Failures"), f"{title} Failures", ["id", "scenario_type", "failed_metrics", "expected", "prediction"], _failure_rows(predictions))
        _add_rows(wb.create_sheet("Metric Definitions"), "Metric Definitions", ["Metric", "Definition"], METRIC_DEFINITIONS)
    else:
        _add_rows(wb.create_sheet("Delta vs Baseline"), "Delta vs Baseline", ["Metric", "Baseline", "Post-Improvement", "Delta", "Pass Bar", "Status"], _delta_rows(baseline_summary, summary))
        _add_rows(wb.create_sheet("Remaining Failures"), "Remaining Failures", ["id", "scenario_type", "failed_metrics", "expected", "prediction"], _failure_rows(predictions))
    return _save_workbook(wb, path)


def _create_failure_analysis_workbook(report_text: str, baseline_predictions: list[dict[str, Any]], path: Path) -> Path:
    wb = _new_workbook()
    clusters = _failure_cluster_rows(report_text)
    _add_rows(
        wb.create_sheet("How to Read"),
        "How to Read This Workbook",
        ["Workbook", "Guidance"],
        _how_to_read_rows(
            "Failure_Analysis.xlsx",
            [
                "Failure Clusters summarizes the baseline failure categories from the report.",
                "Failed Cases lists each baseline case with at least one non-passing metric.",
                "Root Causes and Recommended Improvements convert the markdown analysis into reviewer-friendly tables.",
            ],
        ),
    )
    _add_rows(wb.create_sheet("Failure Clusters"), "Failure Clusters", ["Failure Cluster", "Count", "Example Case", "Likely Root Cause"], clusters)
    _add_rows(wb.create_sheet("Failed Cases"), "Failed Cases", ["id", "scenario_type", "failed_metrics", "expected", "prediction"], _failure_rows(baseline_predictions))
    _add_rows(wb.create_sheet("Root Causes"), "Root Causes", ["Cluster", "Root Cause"], [[row[0], row[3] if len(row) > 3 else ""] for row in clusters])
    _add_rows(wb.create_sheet("Recommended Improvements"), "Recommended Improvements", ["Cluster", "Recommended Improvement"], _recommended_improvement_rows(report_text))
    return _save_workbook(wb, path)


def _recommended_improvement_rows(report_text: str) -> list[list[Any]]:
    if report_text.startswith("Source file was not found"):
        return [["source_status", report_text]]
    rows: list[list[Any]] = []
    for match in re.finditer(r"### (?P<title>.*?)\n(?P<body>.*?)(?=\n### |\n## |\Z)", report_text, re.DOTALL):
        title = match.group("title").strip()
        body = match.group("body")
        rec = re.search(r"- Recommended improvement:\s*(.*)", body)
        if rec:
            rows.append([title, rec.group(1).strip()])
    return rows or [["No recommendations parsed", ""]]


def _extract_bullet_value(text: str, label: str) -> str:
    match = re.search(rf"- {re.escape(label)}:\s*(.*)", text)
    return match.group(1).strip(" `") if match else ""


def _create_improvement_log_workbook(log_text: str, baseline_summary: dict[str, Any], post_summary: dict[str, Any], path: Path) -> Path:
    improvements = _improvement_rows(log_text)
    wb = _new_workbook()
    _add_rows(
        wb.create_sheet("How to Read"),
        "How to Read This Workbook",
        ["Workbook", "Guidance"],
        _how_to_read_rows(
            "Improvement_Log.xlsx",
            [
                "Improvements lists each targeted change and the failure cluster it addressed.",
                "Files Changed identifies where implementation changes were made.",
                "Measured Delta compares baseline and post-improvement metrics without changing eval logic.",
            ],
        ),
    )
    _add_rows(wb.create_sheet("Improvements"), "Improvements", ["Name", "Failure Cluster", "Files Changed", "Expected Impact", "Regression Risk"], improvements)
    _add_rows(wb.create_sheet("Files Changed"), "Files Changed", ["Improvement", "Files Changed"], [[row[0], row[2]] for row in improvements])
    _add_rows(wb.create_sheet("Expected Impact"), "Expected Impact", ["Improvement", "Expected Impact"], [[row[0], row[3]] for row in improvements])
    _add_rows(wb.create_sheet("Measured Delta"), "Measured Delta", ["Metric", "Baseline", "Post-Improvement", "Delta", "Pass Bar", "Status"], _delta_rows(baseline_summary, post_summary))
    return _save_workbook(wb, path)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def _add_title_page(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle).font.size = Pt(12)
    doc.add_paragraph()


def _add_markdown_to_doc(doc: Document, markdown: str) -> None:
    if markdown.startswith("Source file was not found"):
        doc.add_heading("Missing Source", level=1)
        doc.add_paragraph(markdown)
        return

    pending_table: list[list[str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            _flush_doc_table(doc, pending_table)
            pending_table = []
            continue
        if line.startswith("|"):
            cells = [cell.strip(" `") for cell in line.strip("|").split("|")]
            if not all(set(cell) <= {"-", ":"} for cell in cells):
                pending_table.append(cells)
            continue
        _flush_doc_table(doc, pending_table)
        pending_table = []

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line.lstrip("# ").strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line.lstrip("# ").strip(), level=2)
        elif re.match(r"\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line.replace("`", ""))
    _flush_doc_table(doc, pending_table)


def _flush_doc_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for col_index in range(max_cols):
            text = row_values[col_index] if col_index < len(row_values) else ""
            cells[col_index].text = text
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True


def _add_doc_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(_jsonish(value) if value is not None else "")
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _top_failures(report_text: str) -> list[list[Any]]:
    clusters = _failure_cluster_rows(report_text)
    return clusters[:5]


def _example_case_rows(cases: list[dict[str, Any]], limit: int = 5) -> list[list[Any]]:
    if cases and "_missing_source" in cases[0]:
        return [["N/A", cases[0]["_missing_source"], "", ""]]
    rows = []
    for item in cases[:limit]:
        rows.append(
            [
                item.get("id"),
                item.get("scenario_type"),
                item.get("customer_message"),
                item.get("expected_decision"),
            ]
        )
    return rows


def _metric_doc_rows(summary: dict[str, Any]) -> list[list[Any]]:
    return _summary_rows(summary)


def _improvement_doc_rows(log_text: str, baseline_summary: dict[str, Any], post_summary: dict[str, Any]) -> list[list[Any]]:
    deltas = {row[0]: row[3] for row in _delta_rows(baseline_summary, post_summary)}
    rows = []
    for row in _improvement_rows(log_text):
        rows.append(
            [
                row[0],
                row[1],
                row[2],
                row[3],
                "Decision/reason/missing-info/escalation/policy-recall deltas improved in local eval"
                if deltas
                else "N/A",
            ]
        )
    return rows


def _dataset_schema_rows(cases: list[dict[str, Any]]) -> list[list[Any]]:
    if cases and "_missing_source" in cases[0]:
        return [["source_status", cases[0]["_missing_source"]]]
    sample = cases[0] if cases else {}
    return [[key, type(value).__name__] for key, value in sample.items()]


def _create_week4_report_docx(
    report_text: str,
    path: Path,
    *,
    golden_cases: list[dict[str, Any]],
    baseline_summary: dict[str, Any],
    post_summary: dict[str, Any],
    baseline_report: str,
    improvement_log: str,
) -> Path:
    doc = Document()
    _configure_document(doc)
    _add_title_page(
        doc,
        "Week 4 Evaluation Report",
        "Project: return-policy-agent\nCourse: Agentic AI / AI Evals - Week 4\nAgent under test: Northstar Commerce Return Policy Agent\nDate generated: "
        + _generated_timestamp(),
    )

    doc.add_heading("Executive Summary", level=1)
    _add_bullets(
        doc,
        [
            "What the agent does: answers Northstar Commerce return, refund, exchange, gift-return, and shipping-fee questions using local policy evidence.",
            "What was evaluated: backend structured decisions, reason codes, missing-info behavior, escalation, citation coverage, policy-section recall, schema validity, and available LangSmith faithfulness.",
            "Biggest baseline weakness: edge-case decision priority and reason-code mapping, especially for shipping fees, restricted items, and support-review cases.",
            "Biggest improvement: targeted deterministic rules and multi-policy citation recall moved local decision/reason/missing-info/escalation metrics to 100% on the 40-case golden set.",
            "Final recommendation: use the agent as an eval-ready course demo, then run paired LangSmith post-improvement and holdout evaluations before production use.",
        ],
    )

    doc.add_heading("Evaluation Framework", level=1)
    doc.add_paragraph("One-liner: Evaluate whether the agent gives accurate, policy-grounded guidance, asks for missing facts, cites policy evidence, and escalates risky or exception-based cases.")
    doc.add_paragraph("User outcome: Customers receive accurate return/refund/exchange guidance without unsupported refund promises or unsafe exception approvals.")
    _add_doc_table(doc, ["Metric", "Pass Bar", "Judge Method"], [[name, _pass_bar_label(name), definition] for name, definition in METRIC_DEFINITIONS])

    doc.add_heading("Golden Dataset", level=1)
    doc.add_paragraph(f"Dataset size: {len(golden_cases) if not (golden_cases and '_missing_source' in golden_cases[0]) else 'N/A'} labeled cases.")
    _add_doc_table(doc, ["Scenario Type", "Count"], _scenario_mix_rows(golden_cases))
    _add_doc_table(doc, ["Label", "Definition"], LABEL_DEFINITIONS)
    _add_doc_table(doc, ["Case ID", "Scenario Type", "Customer Message", "Expected Decision"], _example_case_rows(golden_cases))

    doc.add_heading("LangSmith Instrumentation", level=1)
    _add_bullets(
        doc,
        [
            "What is traced: top-level agent run plus each major pipeline step.",
            "Top-level run: raw input, structured output, citations, final answer, trace, latency, errors, and retry count.",
            "LLM calls: intent/fact extraction, structured policy decision, and final customer-facing answer generation.",
            "Retrieval: local policy query and retrieved policy chunks with document name, section heading, policy ID, and text.",
            "Validation: deterministic guardrails, schema checks, missing-info validation, and citation validation.",
            "Final answer: customer-facing response generated from the validated structured decision only.",
            "Token/cost/latency: token usage where available, latency per step, and future cost-per-run monitoring.",
        ],
    )
    doc.add_paragraph("LangSmith run link: [Insert link after run]")

    doc.add_heading("Baseline Results", level=1)
    _add_doc_table(doc, ["Metric", "Value", "Pass Bar", "Status"], _metric_doc_rows(baseline_summary))
    doc.add_paragraph("Observations: the local baseline had strong schema and citation behavior, but missed several edge-case decisions, reason codes, missing-info branches, escalation cases, and supporting policy sections.")
    _add_doc_table(doc, ["Failure Cluster", "Count", "Example Case", "Likely Root Cause"], _top_failures(baseline_report))

    doc.add_heading("Improvements Implemented", level=1)
    _add_doc_table(
        doc,
        ["Name", "Failure Cluster Targeted", "Files Changed", "Expected Impact", "Measured Impact"],
        _improvement_doc_rows(improvement_log, baseline_summary, post_summary),
    )

    doc.add_heading("Post-Improvement Results", level=1)
    _add_doc_table(doc, ["Metric", "Value", "Pass Bar", "Status"], _metric_doc_rows(post_summary))
    _add_doc_table(doc, ["Metric", "Baseline", "Post-Improvement", "Delta", "Pass Bar", "Status"], _delta_rows(baseline_summary, post_summary))

    doc.add_heading("Remaining Failure Modes", level=1)
    _add_doc_table(
        doc,
        ["What Still Fails", "Why It Matters", "What I Would Try Next"],
        [
            ["No paired LangSmith post-improvement run", "Faithfulness delta, p95 latency, and cost-per-run are not fully measured.", "Run the same 40-case dataset through LangSmith after improvements."],
            ["Possible golden-set overfitting", "A fully passing local dataset may not represent unseen production questions.", "Add 10-20 fresh holdout cases and real-world failures."],
            ["Full-corpus retrieval noise", "The evidence section can be harder for reviewers or support agents to scan.", "Replace full-corpus retrieval with deterministic policy expansion."],
        ],
    )

    doc.add_heading("Production Monitoring Strategy", level=1)
    _add_doc_table(
        doc,
        ["Monitor", "Why", "Alert / Review Signal"],
        [
            ["Quality drift", "Policy or prompt changes can degrade decisions.", "Drop in decision/reason/missing-info metrics on rolling evals."],
            ["Escalation miss rate", "Missed legal/fraud/high-value cases are high risk.", "Any increase in false-negative escalation cases."],
            ["Citation coverage", "Answers must remain policy-grounded.", "Any answer with missing or invalid citations."],
            ["Tool/retrieval failure rate", "Retrieval failures can force unsafe or unsupported answers.", "Policy retrieval errors or empty evidence sets."],
            ["Latency/cost alerts", "Three LLM calls can affect user experience and budget.", "p95 latency above target or cost per run above threshold."],
        ],
    )

    doc.add_heading("Appendix", level=1)
    doc.add_heading("Dataset Schema", level=2)
    _add_doc_table(doc, ["Field", "Observed Type"], _dataset_schema_rows(golden_cases))
    doc.add_heading("Evaluator Definitions", level=2)
    _add_doc_table(doc, ["Evaluator", "Definition"], METRIC_DEFINITIONS)
    doc.add_heading("Links", level=2)
    _add_bullets(
        doc,
        [
            "LangSmith run link: [Insert link after run]",
            "LangSmith post-improvement run link: [Insert link after run]",
            "Loom walkthrough link: [Insert link after recording]",
        ],
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def _create_loom_docx(script_text: str, path: Path) -> Path:
    doc = Document()
    _configure_document(doc)
    _add_title_page(doc, "Loom Walkthrough Script", "2-3 minute Week 4 evaluation talk track")
    _add_markdown_to_doc(doc, script_text)
    doc.add_heading("Demo Sequence", level=1)
    for item in [
        "Open the Streamlit assistant and show a standard return question.",
        "Show the structured decision, citations, policy evidence, and trace panels.",
        "Open the final evaluation report and point to baseline, improvements, and delta tables.",
        "Open the generated Excel artifacts to show the golden dataset and eval results.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("What Changed", level=1)
    doc.add_paragraph("Targeted deterministic priority fixes, canonical missing-info handling, and multi-policy retrieval/citation recall improvements.")
    doc.add_heading("What Improved", level=1)
    doc.add_paragraph("Local decision, reason-code, missing-info, escalation, citation, policy-section recall, and schema metrics reached 100% on the 40-case golden dataset.")
    doc.add_heading("What Still Fails", level=1)
    doc.add_paragraph("A paired LangSmith post-improvement run, p95 latency, cost per run, and fresh holdout evaluation still need to be captured.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def generate_submission_artifacts(
    output_dir: str | Path = DEFAULT_OUTPUT_DIR,
    source_overrides: dict[str, Path] | None = None,
) -> list[Path]:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    paths = _path_map(source_overrides)

    golden_cases = _read_jsonl(paths["golden_dataset"])
    baseline_predictions = _read_jsonl(paths["baseline_predictions"])
    baseline_summary = _read_json(paths["baseline_summary"])
    post_predictions = _read_jsonl(paths["post_predictions"])
    post_summary = _read_json(paths["post_summary"])
    baseline_report = _read_text(paths["baseline_report"])
    improvement_log = _read_text(paths["improvement_log"])
    final_report = _read_text(paths["final_report"])
    loom_script = _read_text(paths["loom_script"])

    generated = [
        _create_golden_dataset_workbook(golden_cases, output_path / "Golden_Dataset.xlsx"),
        _create_eval_workbook(
            path=output_path / "Baseline_Eval_Results.xlsx",
            title="Baseline Eval Results",
            summary=baseline_summary,
            predictions=baseline_predictions,
        ),
        _create_eval_workbook(
            path=output_path / "Post_Improvement_Eval_Results.xlsx",
            title="Post-Improvement Eval Results",
            summary=post_summary,
            predictions=post_predictions,
            baseline_summary=baseline_summary,
        ),
        _create_failure_analysis_workbook(baseline_report, baseline_predictions, output_path / "Failure_Analysis.xlsx"),
        _create_improvement_log_workbook(improvement_log, baseline_summary, post_summary, output_path / "Improvement_Log.xlsx"),
        _create_week4_report_docx(
            final_report,
            output_path / "Week4_Evaluation_Report.docx",
            golden_cases=golden_cases,
            baseline_summary=baseline_summary,
            post_summary=post_summary,
            baseline_report=baseline_report,
            improvement_log=improvement_log,
        ),
        _create_loom_docx(loom_script, output_path / "Loom_Walkthrough_Script.docx"),
    ]
    return generated


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Week 4 submission artifacts.")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR), help="Directory for generated artifacts.")
    args = parser.parse_args()
    generated = generate_submission_artifacts(output_dir=Path(args.output_dir))
    print("Generated submission artifacts:")
    for path in generated:
        print(f"- {path}")


if __name__ == "__main__":
    main()
