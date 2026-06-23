from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from scripts.generate_submission_artifacts import generate_submission_artifacts


def test_generate_submission_artifacts_creates_expected_files(tmp_path):
    output_dir = tmp_path / "submissions"

    generated = generate_submission_artifacts(output_dir=output_dir)

    expected_files = {
        "Golden_Dataset.xlsx",
        "Baseline_Eval_Results.xlsx",
        "Post_Improvement_Eval_Results.xlsx",
        "Failure_Analysis.xlsx",
        "Improvement_Log.xlsx",
        "Week4_Evaluation_Report.docx",
        "Loom_Walkthrough_Script.docx",
    }
    assert {path.name for path in generated} == expected_files
    assert all(path.exists() for path in generated)

    golden_workbook = load_workbook(output_dir / "Golden_Dataset.xlsx")
    assert golden_workbook.sheetnames == [
        "How to Read",
        "Overview",
        "Golden Dataset",
        "Label Definitions",
        "Scenario Mix",
        "Decision Distribution",
        "Reason Distribution",
    ]
    assert golden_workbook["Golden Dataset"].max_row == 43
    assert golden_workbook["Golden Dataset"].freeze_panes == "A4"
    assert golden_workbook["Golden Dataset"].auto_filter.ref is not None
    assert golden_workbook["Golden Dataset"]["A1"].value == "Golden Dataset"
    assert golden_workbook["Golden Dataset"]["A2"].value == "Generated"

    baseline_workbook = load_workbook(output_dir / "Baseline_Eval_Results.xlsx")
    assert baseline_workbook.sheetnames == [
        "How to Read",
        "Summary Metrics",
        "Scenario Breakdown",
        "Case Results",
        "Failures",
        "Metric Definitions",
    ]
    assert baseline_workbook["Case Results"].freeze_panes == "A4"
    assert baseline_workbook["Case Results"].auto_filter.ref is not None
    assert "decision_accuracy_status" in [
        cell.value for cell in baseline_workbook["Case Results"][3]
    ]

    post_workbook = load_workbook(output_dir / "Post_Improvement_Eval_Results.xlsx")
    assert post_workbook.sheetnames == [
        "How to Read",
        "Summary Metrics",
        "Scenario Breakdown",
        "Case Results",
        "Delta vs Baseline",
        "Remaining Failures",
    ]
    assert "Pass Bar" in [cell.value for cell in post_workbook["Delta vs Baseline"][3]]

    report = Document(output_dir / "Week4_Evaluation_Report.docx")
    paragraphs = [paragraph.text for paragraph in report.paragraphs]
    assert "Week 4 Evaluation Report" in paragraphs
    for expected_heading in [
        "Executive Summary",
        "Evaluation Framework",
        "Golden Dataset",
        "LangSmith Instrumentation",
        "Baseline Results",
        "Improvements Implemented",
        "Post-Improvement Results",
        "Remaining Failure Modes",
        "Production Monitoring Strategy",
        "Appendix",
    ]:
        assert expected_heading in paragraphs
    assert any("LangSmith run link: [Insert link after run]" in text for text in paragraphs)
    assert len(report.tables) >= 8


def test_missing_source_file_adds_placeholder(tmp_path):
    missing = tmp_path / "does-not-exist.jsonl"
    output_dir = tmp_path / "submissions"

    generated = generate_submission_artifacts(
        output_dir=output_dir,
        source_overrides={"golden_dataset": missing},
    )

    assert output_dir / "Golden_Dataset.xlsx" in generated
    workbook = load_workbook(output_dir / "Golden_Dataset.xlsx", read_only=True)
    overview_values = [
        cell.value
        for row in workbook["Overview"].iter_rows()
        for cell in row
        if isinstance(cell.value, str)
    ]
    assert any("Source file was not found" in value for value in overview_values)
